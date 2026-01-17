from docx import Document
from typing import Dict, Any
from utils.logger import Logger

class TemplateAnalyzer:
    """Word模板分析器"""
    
    def __init__(self):
        """初始化分析器"""
        self.logger = Logger('template_analyzer')
    
    def clean_text(self, text: str) -> str:
        """清理文本内容"""
        if not text:
            return ""
        # 移除多余的空白字符
        text = ' '.join(text.split())
        # 移除冒号
        text = text.replace('：', '').replace(':', '')
        return text.strip()
    
    def get_field_value(self, table, field_info):
        """获取指定位置的字段值"""
        try:
            cell = table.cell(field_info['row'], field_info['column'])
            return self.clean_text(cell.text)
        except Exception as e:
            self.logger.error(f"获取字段值失败: {str(e)}")
            return ""
    
    def analyze(self, file_path: str) -> Dict:
        """分析文档并提取信息"""
        try:
            doc = Document(file_path)
            fields = {}
            
            # 获取表格
            if len(doc.tables) > 0:
                table = doc.tables[0]
                
                # 提取字段
                try:
                    # 第一行
                    fields['姓名'] = self.clean_text(table.cell(0, 1).text)  # 第1行第2个单元格
                    fields['性别'] = self.clean_text(table.cell(0, 3).text)  # 第1行第4个单元格
                    fields['出生年月'] = self.clean_text(table.cell(0, 5).text)  # 第1行第6个单元格
                    
                    # 第二行
                    fields['入职日期'] = self.clean_text(table.cell(1, 1).text)  # 第2行第2个单元格
                    fields['文化程度'] = self.clean_text(table.cell(1, 3).text)  # 第2行第4个单元格
                    fields['专业'] = self.clean_text(table.cell(1, 5).text)  # 第2行第6个单元格
                    
                    # 第三行
                    fields['岗位'] = self.clean_text(table.cell(2, 1).text)  # 第3行第2个单元格
                    fields['职称'] = self.clean_text(table.cell(2, 3).text)  # 第3行第4个单元格
                    fields['现从事专业及年限'] = self.clean_text(table.cell(2, 5).text)  # 第3行第6个单元格
                    
                    # 保持原有的提取位置
                    try:
                        # 获取"人员现已具备的条件"的整个单元格内容（在第二列）
                        conditions_cell = None
                        for i, row in enumerate(table.rows):
                            if len(row.cells) > 1 and "人员现已具备的条件" in row.cells[0].text:
                                conditions_cell = row.cells[1]  # 获取第二列的单元格
                                break
                        
                        if conditions_cell:
                            fields['人员现已具备的条件'] = self.clean_text(conditions_cell.text)
                        else:
                            fields['人员现已具备的条件'] = ""
                            
                        # 获取"考核意见"的整个单元格内容（在第二列）
                        opinion_cell = None
                        for i, row in enumerate(table.rows):
                            if len(row.cells) > 1 and "考核意见" in row.cells[0].text:
                                opinion_cell = row.cells[1]  # 获取第二列的单元格
                                break
                        
                        if opinion_cell:
                            fields['考核意见'] = self.clean_text(opinion_cell.text)
                        else:
                            fields['考核意见'] = ""
                            
                    except Exception as e:
                        self.logger.error(f"提取特殊字段失败 {file_path}: {str(e)}")
                        fields['人员现已具备的条件'] = ""
                        fields['考核意见'] = ""
                    
                except IndexError as e:
                    self.logger.error(f"表格结构不符合预期 {file_path}: {str(e)}")
                    return {'fields': {}}
                    
            return {'fields': fields}
            
        except Exception as e:
            self.logger.error(f"分析文档失败 {file_path}: {str(e)}")
            return {'fields': {}} 