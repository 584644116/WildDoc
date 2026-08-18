import os
import re
from datetime import datetime, date
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import pandas as pd
from utils.logger import Logger


class DocumentGenerator:
    """Word文档生成器"""

    def __init__(self):
        self.logger = Logger("document_generator")
        self.template_path = None
        self.template = None
        self.placeholders = {}
        # 映射: 模板占位符 -> Excel 列名，如果不设置则默认为同名映射
        self.field_mapping = {}
        # 命名规则配置
        self.naming_rule = {
            "columns": [],  # 需要作为文件名的列
            "separator": "",  # 分隔符
            "fixed_text": "",  # 追加固定文本
        }

    def load_template(self, template_path: str) -> bool:
        """加载Word模板"""
        try:
            self.template_path = template_path
            self.template = Document(template_path)
            self.placeholders = self._find_placeholders()
            return True
        except Exception as e:
            self.logger.error(f"加载模板失败: {str(e)}")
            return False

    def _find_placeholders(self) -> dict:
        """查找模板中所有的占位符"""
        placeholders = {}
        placeholder_re = re.compile(r"\[(.*?)\]")

        def _extract_text_from_paragraph(paragraph):
            text = paragraph.text or ""
            try:
                elem = paragraph._element
                ns = elem.nsmap or {}
                if "w" not in ns:
                    ns = {
                        **ns,
                        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
                    }
                # Include text from textboxes or other nested runs.
                text_nodes = elem.xpath(".//w:t", namespaces=ns)
                if text_nodes:
                    text = "".join(t.text for t in text_nodes if t.text)
            except Exception:
                # Fallback to paragraph.text if low-level XML access fails.
                pass
            return text

        def _scan_paragraphs(paragraphs, table_index=None, row=None, cell=None):
            for paragraph in paragraphs:
                text = _extract_text_from_paragraph(paragraph)
                if "[" in text and "]" in text:
                    # 占位符可能跨多个run，如"甲[姓名]乙[部门]丙"
                    for placeholder in placeholder_re.findall(text):
                        if not placeholder:
                            continue
                        placeholders[placeholder] = {
                            "table_index": table_index,
                            "row": row,
                            "cell": cell,
                            "paragraph": paragraph,
                            "text": text,
                        }

        # 1) 扫描表格中的占位符
        for table_index, table in enumerate(self.template.tables):
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    _scan_paragraphs(
                        cell.paragraphs, table_index=table_index, row=i, cell=j
                    )

        # 2) 扫描正文段落中的占位符
        _scan_paragraphs(self.template.paragraphs)

        # 3) 扫描页眉/页脚中的占位符
        for section in self.template.sections:
            for header in [
                section.header,
                section.first_page_header,
                section.even_page_header,
            ]:
                _scan_paragraphs(header.paragraphs)
                for t_index, table in enumerate(header.tables):
                    for i, row in enumerate(table.rows):
                        for j, cell in enumerate(row.cells):
                            _scan_paragraphs(
                                cell.paragraphs,
                                table_index=f"header-{t_index}",
                                row=i,
                                cell=j,
                            )
            for footer in [
                section.footer,
                section.first_page_footer,
                section.even_page_footer,
            ]:
                _scan_paragraphs(footer.paragraphs)
                for t_index, table in enumerate(footer.tables):
                    for i, row in enumerate(table.rows):
                        for j, cell in enumerate(row.cells):
                            _scan_paragraphs(
                                cell.paragraphs,
                                table_index=f"footer-{t_index}",
                                row=i,
                                cell=j,
                            )

        return placeholders

    def set_mapping(self, mapping: dict):
        """设置占位符与Excel列的映射"""
        self.field_mapping = mapping or {}

    def set_naming_rule(self, rule: dict):
        """设置文件命名规则"""
        if rule:
            self.naming_rule.update(rule)

    def _format_value(self, value):
        """格式化值，特别处理日期类型"""
        if pd.isna(value):
            return ""
        elif isinstance(value, (datetime, pd.Timestamp)):
            # 如果是datetime类型，只返回日期部分
            return value.strftime("%Y-%m-%d")
        elif isinstance(value, date):
            # 如果是date类型，直接格式化
            return value.strftime("%Y-%m-%d")
        else:
            # 其他类型转换为字符串
            return str(value)

    def _apply_font(self, doc: Document):
        """统一文档字体为宋体 小五"""

        def _set_runs(paragraph):
            for run in paragraph.runs:
                run.font.name = "宋体"
                # East Asia 字体需要额外设置
                try:
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                except Exception:
                    pass
                run.font.size = Pt(9)

        for paragraph in doc.paragraphs:
            _set_runs(paragraph)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _set_runs(paragraph)

    def _generate_file_name(self, row: pd.Series) -> str:
        """根据命名规则生成文件名（不含扩展名）"""
        pattern = self.naming_rule.get("pattern")

        if pattern:
            # 查找所有 [column] 占位符
            placeholders = re.findall(r"\[(.*?)\]", pattern)

            file_name = pattern
            for col in placeholders:
                raw_value = row.get(col, "")
                formatted_value = self._format_value(raw_value)
                file_name = file_name.replace(f"[{col}]", formatted_value)

            base_name = file_name

        else:
            # 回退到旧逻辑或一个默认值
            raw_value = row.get("姓名", "unknown_file")
            base_name = self._format_value(raw_value)

        # 清理非法路径字符
        return re.sub(r'[\\/:*?"<>|]', "_", base_name)

    def generate_single_document(self, row: pd.Series, output_root_dir: str) -> tuple:
        """生成单个文档（线程安全版本）"""
        try:
            # 从模板创建新文档
            doc = Document(self.template_path)

            # 只替换占位符文本，并尽量保留原有“运行(run)”的字体和字号
            # 策略：
            # 1) 优先在单个run内部替换（100%保留该run的字体/字号/样式）；
            # 2) 如占位符被拆分到多个run，退化为将整个段落文本合并到首个run（保留首run样式，避免整段样式丢失）。

            # 构造占位符->Excel列的有效映射（优先用户配置）
            effective_mapping = {
                p: self.field_mapping.get(p, p) for p in self.placeholders.keys()
            }

            def replace_in_paragraph(paragraph):
                replaced_any = False
                # 第一轮：run 内部替换
                for run in paragraph.runs:
                    original = run.text
                    new_text = original
                    for placeholder_key, col_name in effective_mapping.items():
                        token = f"[{placeholder_key}]"
                        if token in new_text:
                            raw_value = row.get(col_name, "")
                            formatted_value = self._format_value(raw_value)
                            new_text = new_text.replace(token, formatted_value)
                    if new_text != original:
                        run.text = new_text
                        replaced_any = True

                # 第二轮：处理跨 run 的占位符（合并到首个 run）
                if not replaced_any and paragraph.runs:
                    combined = "".join(r.text for r in paragraph.runs)
                    original_combined = combined
                    for placeholder_key, col_name in effective_mapping.items():
                        token = f"[{placeholder_key}]"
                        if token in combined:
                            raw_value = row.get(col_name, "")
                            formatted_value = self._format_value(raw_value)
                            combined = combined.replace(token, formatted_value)
                    if combined != original_combined:
                        paragraph.runs[0].text = combined
                        for r in paragraph.runs[1:]:
                            r.text = ""
                        replaced_any = True
                return replaced_any

            # 先处理表格中的段落
            for table in doc.tables:
                for i, table_row in enumerate(table.rows):
                    for j, cell in enumerate(table_row.cells):
                        for paragraph in cell.paragraphs:
                            replace_in_paragraph(paragraph)

            # 再处理文档中不在表格内的段落（以防模板里存在）
            for paragraph in doc.paragraphs:
                replace_in_paragraph(paragraph)

            # 保持模板原有格式：不统一字体
            # self._apply_font(doc)

            # 生成文件名
            file_base_name = self._generate_file_name(row)
            output_file = os.path.join(output_root_dir, f"{file_base_name}.docx")

            # 保存文档
            doc.save(output_file)
            return True, output_file

        except Exception as e:
            return False, str(e)

    def generate_documents(
        self,
        df: pd.DataFrame,
        output_root_dir: str,
        mapping: dict = None,
        naming_rule: dict = None,
        progress_callback=None,
    ) -> None:
        """根据Excel的DataFrame生成Word文档（兼容性方法）"""
        if not self.template_path:
            raise ValueError("请先加载模板文件")

        # 应用外部映射/命名规则（如提供）
        if mapping is not None:
            self.set_mapping(mapping)
        if naming_rule is not None:
            self.set_naming_rule(naming_rule)

        total_rows = len(df)
        # 为每一行数据生成文档
        for index, (row_idx, row) in enumerate(df.iterrows()):
            success, result = self.generate_single_document(row, output_root_dir)
            if success:
                self.logger.info(f"已生成文档: {result}")
            else:
                self.logger.error(f"生成单个文档失败: {result}")

            if progress_callback:
                progress_callback.emit(int((index + 1) * 100 / total_rows))
