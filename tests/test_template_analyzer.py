import unittest
from pathlib import Path
import os
from docx import Document
from core.template_analyzer import TemplateAnalyzer

class TestTemplateAnalyzer(unittest.TestCase):
    def setUp(self):
        """测试前准备"""
        self.test_dir = Path('tests/test_files')
        self.test_dir.mkdir(exist_ok=True)
        
        # 创建测试模板
        self.template_path = self.test_dir / 'test_template.docx'
        self.create_test_template()
        
        self.analyzer = TemplateAnalyzer()
    
    def create_test_template(self):
        """创建测试用Word模板"""
        doc = Document()
        
        # 添加段落
        doc.add_paragraph('这是一个测试模板，包含占位符[姓名]和[年龄]')
        
        # 添加表格
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = '姓名'
        table.cell(0, 1).text = '[姓名]'
        table.cell(1, 0).text = '年龄'
        table.cell(1, 1).text = '[年龄]'
        
        doc.save(self.template_path)
    
    def test_analyze_template(self):
        """测试模板分析"""
        result = self.analyzer.analyze(self.template_path)
        
        # 验证结果
        self.assertIn('placeholders', result)
        self.assertIn('tables', result)
        
        # 验证占位符
        self.assertEqual(set(result['placeholders']), {'姓名', '年龄'})
        
        # 验证表格结构
        self.assertEqual(len(result['tables']), 1)
        table = result['tables'][0]
        self.assertEqual(table['rows'], 2)
        self.assertEqual(table['columns'], 2)
        
        # 验证单元格
        cells = table['structure']['cells']
        self.assertTrue(any(
            cell['text'] == '[姓名]' and cell['has_placeholder']
            for cell in cells
        ))
        self.assertTrue(any(
            cell['text'] == '[年龄]' and cell['has_placeholder']
            for cell in cells
        ))
    
    def tearDown(self):
        """测试后清理"""
        if self.template_path.exists():
            self.template_path.unlink()
        if self.test_dir.exists():
            self.test_dir.rmdir() 