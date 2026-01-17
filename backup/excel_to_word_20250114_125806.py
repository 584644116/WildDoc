import os
import pandas as pd
from docx import Document
from pathlib import Path
from PyQt5.QtWidgets import (QApplication, QMainWindow, QPushButton, QFileDialog, 
                            QLabel, QVBoxLayout, QHBoxLayout, QWidget, QProgressBar, 
                            QTextEdit)
from PyQt5.QtCore import Qt, QThread, pyqtSignal
from PyQt5.QtGui import QFont, QIcon
from docx.enum.text import WD_ALIGN_PARAGRAPH

class GenerateWorkerThread(QThread):
    """工作线程，用于生成Word文档"""
    progress = pyqtSignal(int)  # 进度信号
    log = pyqtSignal(str)      # 日志信号
    finished = pyqtSignal()    # 完成信号
    
    def __init__(self, excel_file, template_file, output_dir):
        super().__init__()
        self.excel_file = excel_file
        self.template_file = template_file
        self.output_dir = output_dir
        
    def generate_word_doc(self, template_path, data, output_path):
        """生成单个Word文档"""
        try:
            # 加载模板
            doc = Document(template_path)
            
            # 替换所有段落中的占位符
            for paragraph in doc.paragraphs:
                for key, value in data.items():
                    placeholder = f"[{key}]"
                    if placeholder in paragraph.text:
                        # 保存段落的对齐方式
                        original_alignment = paragraph.alignment
                        # 替换文本
                        paragraph.text = paragraph.text.replace(placeholder, str(value))
                        # 恢复段落的对齐方式
                        paragraph.alignment = original_alignment
            
            # 处理表格中的占位符
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for key, value in data.items():
                                placeholder = f"[{key}]"
                                if placeholder in paragraph.text:
                                    # 保存段落的对齐方式
                                    original_alignment = paragraph.alignment
                                    # 替换文本
                                    paragraph.text = paragraph.text.replace(placeholder, str(value))
                                    # 恢复段落的对齐方式，保持模板原有格式
                                    paragraph.alignment = original_alignment
                                    # 注释掉强制居中对齐，保持模板原有格式
                                    # paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 保存文档
            doc.save(output_path)
            return True
            
        except Exception as e:
            self.log.emit(f"错误: 生成文档时出错: {str(e)}")
            return False
    
    def run(self):
        """执行文档生成"""
        try:
            # 读取Excel文件
            df = pd.read_excel(self.excel_file)
            total_rows = len(df)
            
            # 创建输出目录
            os.makedirs(self.output_dir, exist_ok=True)
            
            # 处理每一行数据
            for index, row in df.iterrows():
                # 转换为字典
                data = row.to_dict()
                
                # 生成输出文件名
                output_filename = f"{data['姓名']}人员岗位确认表.docx"
                output_path = os.path.join(self.output_dir, output_filename)
                
                # 生成文档
                if self.generate_word_doc(self.template_file, data, output_path):
                    self.log.emit(f"成功生成文档: {output_filename}")
                else:
                    self.log.emit(f"生成文档失败: {output_filename}")
                
                # 更新进度
                progress = int((index + 1) / total_rows * 100)
                self.progress.emit(progress)
            
            self.log.emit("\n所有文档生成完成！")
            self.finished.emit()
            
        except Exception as e:
            self.log.emit(f"错误: {str(e)}")
            self.finished.emit()

class ExcelToWordWindow(QMainWindow):
    """Excel转Word主窗口"""
    def __init__(self):
        super().__init__()
        self.initUI()
        
    def initUI(self):
        """初始化UI"""
        # 设置窗口
        self.setWindowTitle('Word文档生成器')
        self.setGeometry(100, 100, 800, 600)
        self.setStyleSheet("""
            QMainWindow {
                background-color: #f0f2f5;
            }
            QPushButton {
                padding: 8px;
                border-radius: 4px;
                min-width: 100px;
            }
            QProgressBar {
                border: 1px solid #bbb;
                border-radius: 3px;
                text-align: center;
            }
            QProgressBar::chunk {
                background-color: #3498db;
            }
            QTextEdit {
                border: 1px solid #bbb;
                border-radius: 3px;
                padding: 5px;
                background: white;
            }
        """)
        
        # 创建中心部件和布局
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        layout = QVBoxLayout(central_widget)
        
        # 创建按钮
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(10)
        
        self.excel_btn = QPushButton('选择Excel文件', self)
        self.excel_btn.clicked.connect(self.select_excel_file)
        self.excel_btn.setStyleSheet("background-color: #2ecc71; color: white;")
        
        self.template_btn = QPushButton('选择Word模板', self)
        self.template_btn.clicked.connect(self.select_template_file)
        self.template_btn.setStyleSheet("background-color: #3498db; color: white;")
        
        self.output_btn = QPushButton('选择输出目录', self)
        self.output_btn.clicked.connect(self.select_output_dir)
        self.output_btn.setStyleSheet("background-color: #e74c3c; color: white;")
        
        self.generate_btn = QPushButton('生成文档', self)
        self.generate_btn.clicked.connect(self.start_generation)
        self.generate_btn.setEnabled(False)
        self.generate_btn.setStyleSheet("""
            QPushButton {
                background-color: #9b59b6;
                color: white;
            }
            QPushButton:disabled {
                background-color: #bdc3c7;
            }
        """)
        
        # 添加返回按钮
        self.switch_btn = QPushButton('返回提取模式', self)
        self.switch_btn.setStyleSheet("""
            QPushButton {
                background-color: #34495e;
                color: white;
            }
            QPushButton:hover {
                background-color: #2c3e50;
            }
        """)
        self.switch_btn.clicked.connect(self.hide)
        
        # 将按钮添加到布局
        for btn in [self.excel_btn, self.template_btn, self.output_btn, 
                    self.generate_btn, self.switch_btn]:
            btn_layout.addWidget(btn)
        
        layout.addLayout(btn_layout)
        
        # 创建进度条
        self.progress_bar = QProgressBar(self)
        layout.addWidget(self.progress_bar)
        
        # 创建日志显示区域
        self.log_text = QTextEdit(self)
        self.log_text.setReadOnly(True)
        layout.addWidget(self.log_text)
        
        # 初始化文件路径
        self.excel_file = None
        self.template_file = None
        self.output_dir = None
        
        self.show()
    
    def select_excel_file(self):
        """选择Excel文件"""
        file_name, _ = QFileDialog.getOpenFileName(self, '选择Excel文件', '', 
                                                 'Excel Files (*.xlsx *.xls)')
        if file_name:
            self.excel_file = file_name
            self.log_text.append(f'已选择Excel文件: {file_name}')
            self.check_ready()
    
    def select_template_file(self):
        """选择Word模板文件"""
        file_name, _ = QFileDialog.getOpenFileName(self, '选择Word模板', '', 
                                                 'Word Files (*.docx *.doc)')
        if file_name:
            self.template_file = file_name
            self.log_text.append(f'已选择Word模板: {file_name}')
            self.check_ready()
    
    def select_output_dir(self):
        """选择输出目录"""
        dir_name = QFileDialog.getExistingDirectory(self, '选择输出目录')
        if dir_name:
            self.output_dir = dir_name
            self.log_text.append(f'已选择输出目录: {dir_name}')
            self.check_ready()
    
    def check_ready(self):
        """检查是否可以开始生成"""
        self.generate_btn.setEnabled(
            all([self.excel_file, self.template_file, self.output_dir])
        )
    
    def start_generation(self):
        """开始生成文档"""
        self.generate_btn.setEnabled(False)
        self.progress_bar.setValue(0)
        
        # 创建工作线程
        self.worker = GenerateWorkerThread(
            self.excel_file, 
            self.template_file, 
            self.output_dir
        )
        
        # 连接信号
        self.worker.progress.connect(self.progress_bar.setValue)
        self.worker.log.connect(self.log_text.append)
        self.worker.finished.connect(self.generation_finished)
        
        # 开始处理
        self.worker.start()
    
    def generation_finished(self):
        """生成完成的处理"""
        self.generate_btn.setEnabled(True)
        self.progress_bar.setValue(100)

if __name__ == '__main__':
    import sys
    app = QApplication(sys.argv)
    ex = ExcelToWordWindow()
    sys.exit(app.exec_()) 